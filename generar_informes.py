"""
Script para generar tres informes Word simulando opiniones de estudiantes
de último semestre de Derecho sobre la aplicación Justicia IA.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DEST = os.path.join(os.path.dirname(__file__), "Documentoss")
os.makedirs(DEST, exist_ok=True)

def set_heading(doc, text, level=1, color=RGBColor(0x1A, 0x1A, 0x8C)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_paragraph(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    run2 = p.add_run(text)
    run2.italic = italic
    run2.font.size = Pt(11)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A1A8C')
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_header(doc, nombre, codigo, fecha, asignatura):
    doc.add_picture  # solo si hubiera logo
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    cell_left = t.cell(0, 0)
    cell_left.text = ""
    p = cell_left.paragraphs[0]
    r = p.add_run("UNIVERSIDAD LIBRE DE COLOMBIA\nFACULTAD DE DERECHO\nÚLTIMO SEMESTRE — SEMILLERO DE INNOVACIÓN JUDICIAL")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8C)

    cell_right = t.cell(0, 1)
    p2 = cell_right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"Estudiante: {nombre}\nCódigo: {codigo}\nFecha: {fecha}\nAsignatura: {asignatura}")
    r2.font.size = Pt(9)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────
#  INFORME 1 — Valentina Ospina (Entusiasta, muy positivo)
# ─────────────────────────────────────────────────────────────────
def informe_1():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    build_header(doc, "Valentina Ospina Ríos", "2021-0342", "28 de abril de 2026",
                 "Innovación Tecnológica en el Derecho")

    set_heading(doc, "INFORME DE EXPERIENCIA DE USUARIO", level=1)
    set_heading(doc, "Aplicación: Justicia IA — Asistente Judicial Inteligente", level=2,
                color=RGBColor(0x2E, 0x86, 0xC1))

    add_divider(doc)
    doc.add_paragraph()

    set_heading(doc, "1. Introducción", level=2)
    add_paragraph(doc,
        "Como estudiante de último semestre de Derecho con énfasis en litigación civil, "
        "tuve la oportunidad de explorar durante dos sesiones de práctica la plataforma "
        "Justicia IA, un asistente judicial inteligente desarrollado con tecnología de "
        "Inteligencia Artificial Generativa (LLM) sobre infraestructura Groq. El presente "
        "informe recoge mis impresiones, hallazgos y valoración crítica desde la perspectiva "
        "de una futura funcionaria del despacho judicial.")

    set_heading(doc, "2. Descripción General de la Herramienta", level=2)
    add_paragraph(doc,
        "Justicia IA es una plataforma web construida sobre Streamlit que permite a "
        "funcionarios judiciales cargar expedientes en formato PDF o imagen, extraer "
        "automáticamente el texto, categorizarlos por rama del derecho (Civil, Penal, "
        "Laboral, Familia, Administrativo) y generar, en cuestión de segundos, un borrador "
        "de sentencia y un análisis técnico de las pruebas aportadas. Adicionalmente, "
        "incorpora un chat consultivo que permite formular preguntas directas al expediente.")

    set_heading(doc, "3. Experiencia de Uso", level=2)
    add_paragraph(doc,
        "Desde el primer acceso, la interfaz resultó intuitiva. El flujo de trabajo está "
        "muy bien diseñado en cuatro pasos: seleccionar la rama judicial, elegir una carpeta "
        "de destino, cargar los archivos del expediente y ejecutar el análisis. Este orden "
        "lógico reduce la curva de aprendizaje de forma notable.")
    add_paragraph(doc,
        "Cargué un expediente de proceso civil ordinario de 38 páginas. En menos de 12 "
        "segundos el sistema entregó: (i) un borrador de sentencia estructurado con "
        "considerandos, (ii) un análisis de la fuerza probatoria de cada prueba aportada, "
        "y (iii) la categorización automática del proceso. Esto que normalmente me toma "
        "entre 2 y 3 horas de lectura analítica fue sintetizado de manera coherente.")

    set_heading(doc, "4. Funcionalidades Destacadas", level=2)
    items = [
        ("Gestión por Carpetas:", "Permite organizar los expedientes por año, despacho o tipo de proceso, algo que los funcionarios necesitan urgentemente dado el volumen de trabajo."),
        ("Workspace Judicial 360°:", "La vista de tres columnas (Sentencia | Pruebas | Chat) es brillante: permite trabajar el expediente de forma integral sin cambiar de pantalla."),
        ("Consulta de Jurisprudencia:", "Aunque en versión simulada, la búsqueda de precedentes por base de datos vectorial es un avance enorme. En producción sería un diferenciador clave."),
        ("Auto-Retry con Groq:", "La lógica de reintento ante errores de tasa de la API es un detalle de ingeniería que habla de madurez del producto."),
        ("Temas Visuales:", "El modo 'Oscuro Judicial' para reducir la fatiga visual en jornadas largas es una consideración de ergonomía digital que raramente se ve en herramientas jurídicas."),
    ]
    for bold, text in items:
        add_paragraph(doc, text, bold_prefix=bold)

    set_heading(doc, "5. Impacto en la Vida Real del Despacho Judicial", level=2)
    add_paragraph(doc,
        "Colombia tiene actualmente más de 3 millones de procesos represados en el sistema "
        "judicial. Una herramienta como Justicia IA podría reducir el tiempo de análisis "
        "por expediente en un 70-80%, permitiendo que los jueces y magistrados se concentren "
        "en la argumentación jurídica en lugar de la lectura extensiva de documentos. "
        "Estimo que un despacho con esta herramienta podría despachar el triple de procesos "
        "en el mismo período de tiempo.")
    add_paragraph(doc,
        "Es importante resaltar que el sistema siempre aclara que sus borradores son un "
        "'Asistente de Soporte a la Decisión', preservando la responsabilidad y autonomía "
        "del funcionario judicial. Esta postura ética es fundamental para la adopción "
        "institucional de la herramienta.")

    set_heading(doc, "6. Observaciones y Sugerencias", level=2)
    add_paragraph(doc,
        "Mi única observación es que la integración con la API real de la Rama Judicial "
        "(hoy simulada) sería el gran salto para hacer de esto un producto institucional. "
        "También sugiero incorporar la posibilidad de exportar el borrador de sentencia "
        "directamente en formato Word para que el juez lo pueda editar antes de firmar.")

    set_heading(doc, "7. Conclusión", level=2)
    add_paragraph(doc,
        "Justicia IA es, sin duda, el prototipo más completo y útil que he visto en el "
        "contexto académico aplicado al Derecho colombiano. Tiene potencial real de "
        "implementación institucional. Si se conecta con la base de datos de la Rama "
        "Judicial y se incorpora un motor RAG con la jurisprudencia de la Corte Suprema "
        "y el Consejo de Estado, estaríamos frente a una revolución en la administración "
        "de justicia del país. Calificación personal: 9.5/10.")

    add_divider(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Valentina Ospina Ríos\nEstudiante de Derecho — X Semestre\nUniversidad Libre de Colombia")
    r.italic = True
    r.font.size = Pt(10)

    out = os.path.join(DEST, "Informe_1_Valentina_Ospina.docx")
    doc.save(out)
    print(f"Guardado: {out}")

# ─────────────────────────────────────────────────────────────────
#  INFORME 2 — Sebastián Morales (Crítico constructivo, equilibrado)
# ─────────────────────────────────────────────────────────────────
def informe_2():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    build_header(doc, "Sebastián Morales Gutiérrez", "2020-0189", "28 de abril de 2026",
                 "Derecho Procesal y Nuevas Tecnologías")

    set_heading(doc, "REPORTE CRÍTICO DE EVALUACIÓN TECNOLÓGICA", level=1)
    set_heading(doc, "Herramienta evaluada: Justicia IA v2.0", level=2,
                color=RGBColor(0x6C, 0x35, 0x8A))

    add_divider(doc)
    doc.add_paragraph()

    set_heading(doc, "I. Contexto del Evaluador", level=2)
    add_paragraph(doc,
        "Soy estudiante de último semestre con énfasis en Derecho Procesal Penal y he "
        "tenido la oportunidad de hacer práctica en la Fiscalía General de la Nación "
        "durante el último año. Evalúo esta herramienta desde esa perspectiva: alguien "
        "que conoce de primera mano los cuellos de botella del sistema judicial y que "
        "ha manejado expedientes reales de alta complejidad.")

    set_heading(doc, "II. Aspectos Positivos", level=2)
    add_paragraph(doc,
        "La arquitectura del sistema está bien pensada. El flujo de Carpeta → Archivos → "
        "Análisis es coherente con la forma en que un funcionario judicial organiza su "
        "trabajo. La velocidad de procesamiento usando la infraestructura Groq LPU es "
        "impresionante: documentos extensos procesados en segundos es algo que las "
        "herramientas tradicionales como SIGLO no ofrecen.")
    add_paragraph(doc,
        "El análisis de pruebas es la funcionalidad que más me convenció. En materia penal, "
        "la valoración probatoria es el corazón de cualquier decisión. Ver que la IA puede "
        "sintetizar los hechos detectados y evaluar la fuerza de cada prueba —aunque de "
        "forma orientativa— le da al fiscal o al juez un punto de partida muy valioso.")
    add_paragraph(doc,
        "La funcionalidad de cargar una prueba adicional desde el Workspace y obtener su "
        "interpretación inmediata es especialmente poderosa en contextos donde llega un "
        "documento nuevo al expediente a última hora.")

    set_heading(doc, "III. Observaciones Críticas", level=2)
    add_paragraph(doc,
        "Mi principal preocupación, vista desde el Derecho Procesal Penal, es el riesgo "
        "de las 'alucinaciones' del modelo de lenguaje. La propia plataforma lo advierte, "
        "pero en la práctica he visto cómo operadores del derecho tienden a tomar el "
        "primer borrador como verdad sin verificarlo. En un proceso penal, un dato "
        "inventado podría derivar en una privación injusta de la libertad.",
        bold_prefix="⚠️ Riesgo de alucinaciones:")

    add_paragraph(doc,
        "Los datos se procesan mediante la API de Groq, un servicio externo. En un "
        "despacho judicial real, los expedientes contienen información personal sensible "
        "y secreto sumarial. La plataforma reconoce esto y propone el uso de LLMs locales "
        "(Ollama) en producción, lo cual es la solución correcta, pero que en la versión "
        "actual no está implementada.",
        bold_prefix="🔐 Privacidad de datos:")

    add_paragraph(doc,
        "La búsqueda de jurisprudencia y precedentes hoy es una simulación (MockVectorDB). "
        "Para que sea realmente útil necesita conectarse a bases como el SIRI de la Corte "
        "Constitucional o el buscador de jurisprudencia del Consejo de Estado. Sin datos "
        "reales, el resultado es orientativo pero puede inducir a error.",
        bold_prefix="📚 Base jurisprudencial simulada:")

    set_heading(doc, "IV. Potencial de Implementación Real", level=2)
    add_paragraph(doc,
        "Creo firmemente que herramientas como esta son el futuro del despacho judicial "
        "colombiano. La congestión judicial es una crisis estructural y la IA generativa "
        "es una de las pocas soluciones escalables disponibles hoy. Sin embargo, su "
        "implementación debe ir acompañada de: (1) regulación clara sobre el uso de IA "
        "en decisiones judiciales, (2) capacitación obligatoria para los operadores, "
        "y (3) auditorías periódicas de los sesgos del modelo.")

    set_heading(doc, "V. Sugerencias de Mejora", level=2)
    mejoras = [
        "Implementar un sistema de trazabilidad que registre cada sugerencia de la IA y la compare con la decisión final del juez.",
        "Integrar un módulo de detección de sesgos que alerte cuando el modelo pueda estar tomando decisiones estadísticas basadas en raza, género o condición socioeconómica.",
        "Desarrollar un modo 'sin conexión' con LLMs locales para garantizar el secreto sumarial.",
        "Conectar la base vectorial con jurisprudencia real de las Altas Cortes colombianas.",
    ]
    for m in mejoras:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(m).font.size = Pt(11)

    set_heading(doc, "VI. Conclusión", level=2)
    add_paragraph(doc,
        "Justicia IA es un prototipo serio y bien ejecutado que identifica correctamente "
        "el problema de la congestión judicial y propone una solución tecnológica viable. "
        "Mis observaciones no son para desmeritar el trabajo sino para señalar los riesgos "
        "que deben resolverse antes de una implementación institucional. Con los ajustes "
        "descritos, podría convertirse en una herramienta de referencia en América Latina. "
        "Calificación: 8/10, con alto potencial de mejora.")

    add_divider(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Sebastián Morales Gutiérrez\nEstudiante de Derecho — X Semestre\nÉnfasis: Derecho Procesal Penal")
    r.italic = True
    r.font.size = Pt(10)

    out = os.path.join(DEST, "Informe_2_Sebastian_Morales.docx")
    doc.save(out)
    print(f"Guardado: {out}")

# ─────────────────────────────────────────────────────────────────
#  INFORME 3 — Daniela Cárdenas (Escéptica, enfoque ético/social)
# ─────────────────────────────────────────────────────────────────
def informe_3():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    build_header(doc, "Daniela Cárdenas Melo", "2020-0567", "28 de abril de 2026",
                 "Derecho Constitucional y Derechos Humanos")

    set_heading(doc, "ANÁLISIS ÉTICO-JURÍDICO DE HERRAMIENTA DE IA JUDICIAL", level=1)
    set_heading(doc, "Justicia IA: ¿Avance o riesgo para el acceso a la justicia?", level=2,
                color=RGBColor(0xC0, 0x39, 0x2B))

    add_divider(doc)
    doc.add_paragraph()

    set_heading(doc, "Presentación", level=2)
    add_paragraph(doc,
        "Escribo este informe desde mi perspectiva como estudiante de Derecho Constitucional "
        "y activista en derechos humanos. He participado en clínicas jurídicas con comunidades "
        "vulnerables y mi mirada sobre la tecnología judicial siempre parte de una pregunta "
        "fundamental: ¿a quién beneficia y a quién puede perjudicar? Evalúo Justicia IA con "
        "ese mismo lente crítico.")

    set_heading(doc, "1. Lo que la herramienta hace bien", level=2)
    add_paragraph(doc,
        "Reconozco que Justicia IA ataca un problema real y documentado: la congestión "
        "judicial en Colombia es una crisis que afecta especialmente a las personas de "
        "menores recursos, que no pueden costear abogados que aceleren sus procesos. "
        "Si esta herramienta ayuda a los despachos a resolver más rápido, el principal "
        "beneficiario debería ser la ciudadanía.")
    add_paragraph(doc,
        "La interfaz es accesible e intuitiva. No requiere conocimientos técnicos profundos, "
        "lo que la hace usable por funcionarios judiciales sin formación en tecnología. "
        "El sistema de temas visuales —incluyendo el modo de alto contraste— muestra "
        "consideración por la diversidad de usuarios.")
    add_paragraph(doc,
        "La advertencia permanente de que 'los borradores generados por la IA no son "
        "vinculantes y la responsabilidad recae en el funcionario humano' es éticamente "
        "necesaria y me alegra que esté presente en todas las interfaces. Muchos sistemas "
        "similares omiten esta salvaguarda.")

    set_heading(doc, "2. Preocupaciones Fundamentales", level=2)

    add_paragraph(doc,
        "Los modelos de lenguaje como Llama 3 se entrenan con corpus de texto en inglés "
        "de países del Norte Global. La jurisprudencia colombiana, el derecho social "
        "latinoamericano y las particularidades de nuestro sistema procesal están "
        "subrepresentados. Esto puede llevar a que el sistema 'proponga' soluciones "
        "jurídicas ajenas a nuestro contexto o que repliquen sesgos históricos contra "
        "comunidades étnicas, mujeres víctimas de violencia o trabajadores informales.",
        bold_prefix="2.1. Sesgos estructurales del modelo:")

    add_paragraph(doc,
        "Tuve la oportunidad de cargar un expediente de un proceso de restitución de "
        "tierras. La categoría asignada fue 'Civil' cuando debería haberse reconocido "
        "como un proceso especial de la Ley 1448. El chat respondió correctamente cuando "
        "pregunté directamente, pero el triage automático falló en la especificidad. "
        "En casos de víctimas del conflicto, estos errores tienen consecuencias humanas.",
        bold_prefix="2.2. Caso práctico observado:")

    add_paragraph(doc,
        "La plataforma funciona con una API externa (Groq). El README reconoce este "
        "riesgo y propone LLMs locales para producción. Pero mientras eso no ocurra, "
        "los expedientes que se procesen —incluyendo datos de víctimas, menores de edad "
        "o personas en situación de vulnerabilidad— estarían pasando por servidores "
        "externos. Desde el Derecho Constitucional, esto puede vulnerar el Habeas Data "
        "y el artículo 15 de la Constitución Política.",
        bold_prefix="2.3. Habeas Data y soberanía de la información:")

    add_paragraph(doc,
        "La promesa de 'reducir el tiempo de análisis' no debe convertirse en una "
        "presión para que los jueces 'despachen más rápido'. La calidad de la justicia "
        "no se mide en velocidad. Un juez que confíe ciegamente en el borrador de la "
        "IA y lo firme sin la reflexión jurídica correspondiente estaría negando el "
        "derecho al debido proceso de las partes.",
        bold_prefix="2.4. Riesgo de delegación acrítica:")

    set_heading(doc, "3. Propuestas desde el Derecho Constitucional", level=2)
    propuestas = [
        "Establecer por norma que ningún borrador de IA puede incorporarse a una providencia judicial sin la revisión y rúbrica explícita del funcionario, con constancia de que se revisó el expediente original.",
        "Crear un comité de auditoría de sesgos con participación de organizaciones de derechos humanos, movimientos de mujeres y representantes de comunidades étnicas.",
        "Exigir que el modelo se entrene o fine-tune con corpus jurídico colombiano antes de su uso institucional.",
        "Desarrollar un protocolo especial para expedientes que involucren víctimas del conflicto armado, menores o población LGBTIQ+.",
        "Garantizar que el modelo sea de código abierto y auditado por el Ministerio de Justicia antes de su despliegue oficial.",
    ]
    for pr in propuestas:
        p = doc.add_paragraph(style='List Number')
        p.add_run(pr).font.size = Pt(11)

    set_heading(doc, "4. Valoración Final", level=2)
    add_paragraph(doc,
        "Justicia IA es una iniciativa valiosa y técnicamente competente. Pero el "
        "problema de la justicia en Colombia no es solo de velocidad: es de acceso, "
        "de equidad y de confianza institucional. La tecnología, por sí sola, no resuelve "
        "la inequidad estructural del sistema. Si se implementa con las salvaguardas "
        "adecuadas, regulación robusta y un enfoque diferencial, puede ser un gran aliado. "
        "Si se implementa apresuradamente, puede reproducir y escalar las mismas "
        "desigualdades que hoy afectan a los más vulnerables.")
    add_paragraph(doc,
        "Mi calificación del prototipo como herramienta tecnológica: 8/10. "
        "Mi calificación de la reflexión ética incluida en el sistema: 7/10 "
        "(hay conciencia del problema, pero aún falta profundidad en las soluciones). "
        "Potencial de impacto social positivo, con las condiciones correctas: muy alto.",
        italic=True)

    add_divider(doc)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Daniela Cárdenas Melo\nEstudiante de Derecho — X Semestre\nÉnfasis: Derecho Constitucional y DDHH")
    r.italic = True
    r.font.size = Pt(10)

    out = os.path.join(DEST, "Informe_3_Daniela_Cardenas.docx")
    doc.save(out)
    print(f"Guardado: {out}")


if __name__ == "__main__":
    informe_1()
    informe_2()
    informe_3()
    print("\nLos tres informes han sido generados en la carpeta 'Documentoss'.")
